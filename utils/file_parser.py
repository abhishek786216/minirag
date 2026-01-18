import PyPDF2
import docx
import io
import aiofiles
from typing import Dict, Optional
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

class FileParser:
    """Parse different file types and extract text content"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.docx'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    @classmethod
    async def parse_file(cls, file_content: bytes, filename: str) -> Dict[str, str]:
        """
        Parse file content and extract text
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            
        Returns:
            Dictionary with text content and metadata
            
        Raises:
            ValueError: If file type not supported or file too large
        """
        file_path = Path(filename)
        extension = file_path.suffix.lower()
        
        if extension not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}. Supported: {cls.SUPPORTED_EXTENSIONS}")
        
        if len(file_content) > cls.MAX_FILE_SIZE:
            raise ValueError(f"File too large: {len(file_content)} bytes. Max: {cls.MAX_FILE_SIZE} bytes")
        
        try:
            if extension == '.pdf':
                text = cls._parse_pdf(file_content)
            elif extension == '.txt':
                text = cls._parse_txt(file_content)
            elif extension == '.docx':
                text = cls._parse_docx(file_content)
            else:
                raise ValueError(f"Unsupported extension: {extension}")
            
            if not text.strip():
                raise ValueError("No text content found in file")
            
            logger.info(f"Parsed {filename}: {len(text)} characters extracted")
            
            return {
                'text': text,
                'filename': filename,
                'file_type': extension[1:],  # Remove dot
                'char_count': len(text),
                'word_count': len(text.split())
            }
            
        except Exception as e:
            logger.error(f"Error parsing {filename}: {str(e)}")
            raise ValueError(f"Failed to parse {filename}: {str(e)}")
    
    @staticmethod
    def _parse_pdf(file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    text += page_text + "\n\n"
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {e}")
                    continue
            
            if not text.strip():
                raise ValueError("No text could be extracted from PDF")
            
            return text.strip()
            
        except Exception as e:
            raise ValueError(f"PDF parsing error: {str(e)}")
    
    @staticmethod
    def _parse_txt(file_content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try UTF-8 first
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                # Fallback to latin-1
                try:
                    text = file_content.decode('latin-1')
                except UnicodeDecodeError:
                    # Final fallback
                    text = file_content.decode('utf-8', errors='ignore')
            
            return text.strip()
            
        except Exception as e:
            raise ValueError(f"TXT parsing error: {str(e)}")
    
    @staticmethod
    def _parse_docx(file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            if not text.strip():
                raise ValueError("No text could be extracted from DOCX")
            
            return text.strip()
            
        except Exception as e:
            raise ValueError(f"DOCX parsing error: {str(e)}")
    
    @classmethod
    def validate_file(cls, filename: str, file_size: int) -> None:
        """
        Validate file before processing
        
        Args:
            filename: Name of the file
            file_size: Size in bytes
            
        Raises:
            ValueError: If validation fails
        """
        file_path = Path(filename)
        extension = file_path.suffix.lower()
        
        if extension not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")
        
        if file_size > cls.MAX_FILE_SIZE:
            raise ValueError(f"File too large: {file_size} bytes")
        
        if file_size == 0:
            raise ValueError("File is empty")